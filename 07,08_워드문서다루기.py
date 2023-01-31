from docx import Document






import requests
from bs4 import BeautifulSoup
import time
import pyautogui

#사용자 입력 파트
keyword = pyautogui.prompt('검색어를 입력하세요')
lastpage = int(pyautogui.prompt('몇 페이지까지 가져올까요'))

#1.워드생성하기
document = Document()
pagenum = 1
for i in range(1, lastpage *10 , 10):
    print(f'{pagenum}페이지 크롤링 중입니다.===============================================================')

    response =requests.get(f'https://search.naver.com/search.naver?sm=tab_hty.top&where=news&query={keyword}&start={i}')
    html = response.text
    soup = BeautifulSoup(html, 'html.parser')

    articles = soup.select('div.info_group')
    for article in articles:
        links = article.select('a.info')
        if len(links) >= 2:
            url = links[1].attrs['href']#링크가 2개이상이면 두 번째 링크의 href를 추출
            response = requests.get(url, headers = {'User-agent':'Mozila/5.0'})
            html = response.text
            soup= BeautifulSoup(html,'html.parser')

            #만약 연예뉴스라면
            if 'entertain' in response.url:
                title = soup.select_one('.end_tit') #기사 제목
                content = soup.select_one('#articeBody')
            elif 'sports' in response.url:
                title = soup.select_one('h4.title') #기사 제목 그냥 .title을 하면 50개가 나오므로 제목만 나오게 h4로 한정시켜줌
                content = soup.select_one('#newsEndContents')
                #본문 내용 안에 불필요한 div, p 삭제하기
                divs = content.select('div')
                for d in divs:
                    d.decompose() #d태그 삭제

                ps = content.select('p')
                for p in ps:
                    p.decompose() #p태그 삭제
        
            else:
                title = soup.select_one('.media_end_head_top')
                content = soup.select_one("#newsct_article") #뉴스 본문의 id값
            print("====링크====\n", url)
            print("====제목====\n", title.text.strip()) #텍스트만 가져오기
            print("====본문====\n", content.text.strip())
             # 2.워드 데이터 추가하기
            document.add_heading(title.text.strip(), level = 0)
            document.add_paragraph(url)
            document.add_paragraph(content.text.strip())
            time.sleep(0.3)
    pagenum += 1
   

#3.워드 저장하기
document.save('test.docx')

